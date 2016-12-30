""" redox - an engine to assist screenplay writing  """
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys
import os
import re
import time
import threading
import urllib
from socket import gaierror
import requests
from requests.packages.urllib3.exceptions import MaxRetryError
from requests.packages.urllib3.exceptions import NewRequestsError
from requests.exceptions import ConnectionError as RequestsError
from more_itertools import unique_everseen
import docx
from textblob import Blobber
from textblob.np_extractors import ConllExtractor
from textblob_aptagger import PerceptronTagger
from bs4 import BeautifulSoup
from GoogleScraper import scrape_with_config
from PIL import Image
from selenium.common.exceptions import WebDriverException
TB = Blobber(pos_tagger=PerceptronTagger(),
             np_extractor=ConllExtractor())

def cleanimages(img_dir):
    """ Delete smaller images """
    for filename in os.listdir(img_dir):
        filepath = os.path.join(img_dir, filename)
        with Image.open(filepath) as ppix:
            fat, tall = ppix.size
        pixelsize = fat*tall
        if pixelsize < 480000:
            os.remove(filepath)

def get_immediate_subdirectories(a_dir):
    """ Get only the immediate subfolders """
    return [os.path.join(PROJECTPATH, chapter)
            for chapter in os.listdir(a_dir)
            if os.path.isdir(os.path.join(a_dir, chapter))]

def printprogress(iteration, total, prefix='',
                  suffix='', decimals=1, pbarlength=100):
    """ progress bar """
    formatstr = "{0:." + str(decimals) + "f}"
    percent = formatstr.format(100 * (iteration / float(total)))
    filledlength = int(round(pbarlength * iteration / float(total)))
    pbar = '█' * filledlength + '-' * (pbarlength - filledlength)
    sys.stdout.write(
        '\r%s |%s| %s%s %s' % (prefix, pbar, percent, '%', suffix))
    if iteration == total:
        sys.stdout.write('\n')
    sys.stdout.flush()

class FetchResource(threading.Thread):
    """stores a web resource  in the target directory """
    def __init__(self, target, urls):
        super().__init__()
        self.target = target
        self.urls = urls
    def run(self):
        for url in self.urls:
            furl = urllib.parse.unquote(url)
            gurl = furl.split("/")[-1][:10]
            pixname = re.sub('[^0-9a-zA-Z]+', '', gurl)
            if "." not in pixname:
                pixname.join(pixname.join(".jpg"))
                with open(os.path.join(
                    self.target, pixname), 'wb') as pix:
                    try:
                        content = requests.get(furl).content
                        pix.write(content)
                        print('[+] {}'.format(pixname))
                    except (RequestsError,
                            NewRequestsError,
                            MaxRetryError,
                            gaierror):
                        print("\nrequests  error for  {}\n".format(furl))

def phrasescraper(aphrase, aprocpath):
    """ Gets images for a phrase and writes to the phrase folder """
    print("\t\tBeginning scrape for {}".format(aphrase))
    target_directory = os.path.join(aprocpath, "images", aphrase)
    if not os.path.isdir(target_directory):
        os.makedirs(target_directory)
    config = {
        'keyword': aphrase,
        'search_engines': ['yandex', 'yahoo'],
        'search_type': 'image',
        'scrape_method': 'selenium',
        'sel_browser' : 'Phantomjs',
        'do_caching': False,
        'database_name' : aphrase}
    try:
        search = scrape_with_config(config)
        image_urls = []
        if search is not None:
            for serp in search.serps:
                image_urls.extend(
                    [link.link for
                     link in serp.links])
            print('\t\t[i] Saving {num} images at "{dir}"'.
                  format(num=len(image_urls), dir=target_directory))
            num_threads = 100
            linkthreads = [FetchResource(target_directory, [])
                           for i in range(num_threads)]
            while image_urls:
                for linkthread in linkthreads:
                    try:
                        linkthread.urls.append(image_urls.pop())
                    except IndexError:
                        break
            imgthreads = [linkthread
                          for linkthread in linkthreads
                          if linkthread.urls]
            for imgthread in imgthreads:
                imgthread.start()
            for imgthread in imgthreads:
                imgthread.join()
            print('\t\tfinished phrase operations for {} at {}'.
                  format(aphrase, time.strftime('%X')))
        else:
            print("\t\t\nNo results for {}\n".format(aphrase))
    except WebDriverException:
        print("\n\tghostdriver exception for {}!\n".format(aphrase))

def frameoperations(aprocpath, someframes):
    """ file writing and printing """
    chapterdoc = docx.Document()
    chapterdoc.add_heading(aprocpath.split("\\")[-1], 0)
    bodytable = chapterdoc.add_table(rows=len(someframes), cols=1)
    progcount = 0
    printprogress(progcount, len(someframes),
                  prefix='Progress:', suffix='Complete\n',
                  pbarlength=len(someframes))
    for frameno, frame in enumerate(someframes):
        frametable = bodytable.cell(
            frameno, 0).add_table(rows=5, cols=1)
        frameinfos = [[]]
        framephrases = [[]]
        infotext = ""
        tickersubjectivity = 1
        tickerpolarity = 0
        subticker = str(frame[1])
        for sentence in frame:
            for phrase in sentence.noun_phrases:
                framephrases = [unique_everseen(
                    framephrases)].extend(phrase)
                opener = urllib.request.build_opener()
                opener.addheaders = [('User-agent', 'Mozilla/5.0')]
                article = urllib.request.quote(phrase)
                resource = opener.open(
                    "http://en.wikipedia.org/wiki/" + article)
                data = resource.read()
                resource.close()
                soup = BeautifulSoup(data, "lxml")
                infomarkup = soup.find('div', id="bodyContent").p
                for info in infomarkup:
                    infotext += "\n"+ info.text
                frameinfos.extend(infotext)
                print("\n\t\t\t\t {} : {}".
                      format(phrase, infotext))
            if sentence.subjectivity < tickersubjectivity:
                if sentence.polarity > tickerpolarity:
                    subticker = str(sentence)
        print("\n\n\t\t\t{}\n\n".format(str(subticker)))
        frametable.cell(0, 0).text = " ".join(
            [str(tch) for tch in frame]) + "\n" + "-"*60
        frametable.cell(1, 0).text = "\n".join(
            [str(frameinfo) for frameinfo in frameinfos])
        frametable.cell(1, 0).add_paragraph(
            "\n".join(framephrases))
        frametable.cell(1, 0).add_paragraph(
            "\n" + "-"*60)
        frametable.cell(3, 0).text = "Ticker Suggestions :\n*2"
        frametable.cell(3, 0).add_paragraph(
            str(subticker) + "\n" + "-"*60)
        frametable.cell(4, 0).text = "_"*103
        time.sleep(0.1)
        progcount += 1
        printprogress(progcount, len(someframes),
                      prefix='Progress:', suffix='Complete',
                      pbarlength=50)
    chapterdoc.save(os.path.join(
        aprocpath, "rawscreenplay.docx"))
    print("\t\tWrote document at {}".format(aprocpath))

def frameify(aprocpath):
    """ Separate frames  """
    if os.path.exists(os.path.join(aprocpath, "script.txt")):
        scfile = os.path.join(aprocpath, 'script.txt')
        with open(scfile, encoding='ascii', mode='r',
                  errors='ignore') as scriptfile:
            ascripttext = TB(scriptfile.read())
            scripttext = ascripttext.correct()
        frames = [scripttext.sentences[x:x+4]
                  for x in range(0, len(scripttext.sentences), 4)]
        print("\n\tCorrected spellings and frameed {}\n\n".
              format(aprocpath.split("\\")[-1]))
        frameoperations(aprocpath, frames)
    else:
        print("\n\tNo phraselist or script forund for {}".
              format(aprocpath.split("\\")[-1]))

def chapterops(chapterpath):
    """ Reads chapter and calls phrase operations"""
    if os.path.exists(os.path.join(
            chapterpath, "phraselist.txt")):
        with open(os.path.join(
            chapterpath, "phraselist.txt"), 'r', encoding='ascii',
                  errors='ignore') as oldphrasefile:
            framephrases = oldphrasefile.readlines()
        for phrase in framephrases:
            newphrase = "".join(
                [smbl for smbl in phrase
                 if smbl.isalnum() or smbl == " "])
            if phrase is not None:
                phrasescraper(newphrase, chapterpath)
            else:
                print("\t\nEmpty phrase found for {}\n".
                      format(chapterpath.split("\\")[-1]))
        for dirname in os.listdir(
                os.path.join(chapterpath, "images")):
            if os.path.isdir(
                    os.path.join(chapterpath, "images", dirname)):
                cleanimages(os.path.join(
                    os.path.join(chapterpath, "images", dirname)))
            else:
                pass
        print("\tremoved smaller images for {}"
              .format(chapterpath.split("\\")[-1]))

    elif os.path.exists(os.path.join(
            chapterpath, "rawscreenplay.docx")):
        print("{} is ready for first review".
              format(chapterpath.split("\\")[-1]))
    elif not os.path.exists(os.path.join(
            chapterpath, "rawscreenplay.docx")):
        frameify(chapterpath)

if __name__ == '__main__':
    PROJECTPATH = "C:\\Users\\nnikh\\Documents\\scrape"
    print("\nBeginning operations at {}\n".
          format(PROJECTPATH))
    CHAPTERS = get_immediate_subdirectories(PROJECTPATH)
    CHAPTERPOOL = [threading.Thread(
        target=chapterops, args=(chapter,))
                   for chapter in CHAPTERS]
    for proc in CHAPTERPOOL:
        proc.start()
    for proc in CHAPTERPOOL:
        proc.join()
